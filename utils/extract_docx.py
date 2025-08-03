import docx
import sys

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        return f"提取文档内容时出错: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        text = extract_text_from_docx(file_path)
        print(text)
    else:
        print("请提供Word文档路径作为参数")